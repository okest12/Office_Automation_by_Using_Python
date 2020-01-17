import re
import docx
import os
import win32com
from win32com.client import Dispatch

def print_by_docx(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            print(run.text)

def replace_by_docx():
    doc = docx.Document(r"C:\01-code\2-2019年企业所得税审核报告及说明范本.docx")
    print_by_docx(doc)
    for para in doc.paragraphs:
        m = re.search('S[0-9]+[A-Z][0-9]+', para.text)
        while m:
            para.text = para.text.replace(m.group(0), "123456")
            m = re.search('S[0-9]+[A-Z][0-9]+', para.text)
    doc.save(r"C:\01-code\2-2019年企业所得税审核报告及说明范本_new1.docx")

def replace_by_docx_run():
    doc = docx.Document(r"C:\01-code\2-2019年企业所得税审核报告及说明范本_new.docx")
    print_by_docx(doc)
    for para in doc.paragraphs:
        for run in para.runs:
            m = re.search('S[0-9]+[A-Z][0-9]+', run.text)
            if m:
                run.text = "12345"
    doc.save(r"C:\01-code\2-2019年企业所得税审核报告及说明范本_new1.docx")

def replace_by_win32_com():
    a = 1

replace_by_docx_run()